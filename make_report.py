# -*- coding: utf-8 -*-
"""Generate a combined Word report for Homework #1 and #2 (Korean)."""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn

BASE = "/projects/sandbox"
DS = os.path.join(BASE, "dataset")

doc = Document()

# ---- global style: Korean-friendly font, compact ----
style = doc.styles["Normal"]
style.font.name = "Malgun Gothic"
style.font.size = Pt(10)
style._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.08

# margins
for sec in doc.sections:
    sec.top_margin = Cm(1.8)
    sec.bottom_margin = Cm(1.8)
    sec.left_margin = Cm(2.0)
    sec.right_margin = Cm(2.0)


def set_kfont(run, size=None, bold=None, color=None):
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color:
        run.font.color.rgb = color


def heading(text, size=13, space_before=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    set_kfont(r, size=size, bold=True, color=RGBColor(0x1F, 0x37, 0x64))
    return p


def subheading(text, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    set_kfont(r, size=size, bold=True)
    return p


def body(text, size=10):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_kfont(r, size=size)
    return p


def bullet(text, size=10):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    set_kfont(r, size=size)
    return p


def two_images(img1, cap1, img2, cap2, w=7.1):
    """Place two images side by side using a table."""
    t = doc.add_table(rows=2, cols=2)
    t.autofit = True
    for col, (img, cap) in enumerate([(img1, cap1), (img2, cap2)]):
        cell = t.cell(0, col)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(img, width=Cm(w))
        cap_cell = t.cell(1, col)
        cp = cap_cell.paragraphs[0]
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cp.add_run(cap)
        set_kfont(r, size=8.5, bold=False, color=RGBColor(0x55, 0x55, 0x55))
    return t


def one_image(img, cap, w=11.0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(img, width=Cm(w))
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cp.add_run(cap)
    set_kfont(r, size=8.5, color=RGBColor(0x55, 0x55, 0x55))


def make_table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = hdr[i].paragraphs[0].add_run(h)
        set_kfont(r, size=9, bold=True)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = cells[i].paragraphs[0].add_run(str(val))
            set_kfont(r, size=9)
    return t


# =====================================================================
# COVER PAGE (excluded from 4-page limit)
# =====================================================================
for _ in range(4):
    doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("전산재료공학 과제 보고서"); set_kfont(r, size=24, bold=True, color=RGBColor(0x1F, 0x37, 0x64))
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Homework #1 & #2"); set_kfont(r, size=16, bold=True)
for _ in range(2):
    doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("#1. 원형 홀을 포함한 평판의 2차원 선형탄성 유한요소해석\n"
              "#2. 압입시험 데이터 기반 탄성계수(E) 예측 인공신경망")
set_kfont(r, size=12)
for _ in range(8):
    doc.add_paragraph()
for label in ["과    목 : ____________________",
              "학    번 : ____________________",
              "이    름 : ____________________",
              "제 출 일 : 2026. __. __."]:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(label); set_kfont(r, size=12)

doc.add_page_break()

# =====================================================================
# HOMEWORK #1
# =====================================================================
heading("Homework #1. 원형 홀을 포함한 평판의 2차원 선형탄성 유한요소해석 (Plane Stress)", size=13, space_before=0)
body("폭 W=2.0 m, 높이 H=1.0 m, 두께 t=0.01 m인 평판(중심 (0,0))에 중심 (0.1, 0), 반지름 r의 "
     "원형 홀을 둔 모델을 평면응력(Plane Stress) 조건에서 해석하였다. 재료는 E=210 GPa, ν=0.30이며, "
     "하단은 완전 고정(u=v=0), 상단은 위쪽으로 Δ=1×10⁻⁴ m의 변위를 강제하였다(변위 구속 조건).")

subheading("1) 예제 코드 8단계의 역할")
bullet("① 입력 파라미터: 형상(W,H,t), 재료물성(E,ν), 홀 위치·크기, 메쉬 간격(h, n_circle), 경계조건(Δ) 등 해석 조건 정의.")
bullet("② 재료물성·요소 정의: 평면응력 구성행렬 D 계산과 CST(상수변형률 삼각형) 요소의 B행렬·면적 계산 함수 정의.")
bullet("③ 모델·메쉬 생성: 격자점 생성 → 홀 내부 점 제거 → 경계·홀 경계점 추가 → Delaunay 삼각분할 → 홀 내부 요소 제거.")
bullet("④ 경계조건 설정: 하단/상단/좌우 경계 절점을 탐색하여 고정 및 변위 부여 대상 절점을 식별.")
bullet("⑤ 강성행렬 구성 및 해 계산: 요소 강성행렬 Kᵉ를 전체 행렬 K로 조립하고, Dirichlet 경계조건을 소거법으로 적용해 K_ff·U_f = −K_fc·U_c 를 풀어 변위 U를 구함.")
bullet("⑥ 후처리(응력·변형률): 각 요소에서 ε=B·u, σ=D·ε 및 von Mises 응력을 계산.")
bullet("⑦ 시각화 함수: 변형 형상과 응력·변형률 컨투어를 그리는 헬퍼 함수.")
bullet("⑧ 실행(Run): 위 단계를 순서대로 호출하여 해석을 수행하고 결과를 출력·시각화.")

subheading("2) 기본 해석 결과 (변위장·응력분포)")
body("상단을 위로 당기는 세로 인장 하중에 대해, 변위장은 하단 고정부에서 0이고 위로 갈수록 증가하여 상단에서 "
     "강제 변위 Δ에 도달하는 전형적 분포를 보였다. 응력은 판 전반에서 비교적 균일하다가 홀 근처에서 급격히 "
     "교란되는데, 이는 하중 전달 경로가 구멍을 우회하면서 가장자리에 몰리기 때문이다.")
two_images(os.path.join(BASE, "vonmises_r_min.png"), "그림 1. von Mises 응력 분포 (작은 구멍, r=0.05)",
           os.path.join(BASE, "vonmises_r_max.png"), "그림 2. von Mises 응력 분포 (큰 구멍, r=0.30)")
body("[결과 해석] 응력집중은 하중 방향과 수직인 홀의 좌·우 측면(x=cx±r, y=0)에서 가장 크게 나타났고, 하중 "
     "방향과 나란한 홀의 상·하부는 오히려 저응력 영역이 되었다. 이는 원공 주위 응력집중 이론과 정확히 일치하는 "
     "거동이다. 기본 조건(r=0.15)에서 홀 가장자리의 최대 응력(von Mises)은 약 51.5 MPa로, 적용 공칭응력"
     "(σ=E·Δ/H=21 MPa)의 약 2.45배에 해당하는 응력집중이 발생하였다. 한편 본 해석은 변위를 강제하는 조건이므로, "
     "선형탄성에서 변위장은 탄성계수 E와 무관하고(E가 강성행렬에서 약분됨) 응력만 E에 선형 비례한다. 따라서 "
     "예컨대 E를 절반으로 줄이면 변형 형상은 동일하고 응력 크기만 절반이 되며, 응력집중의 정도(SCF)는 변하지 "
     "않는다. 즉 응력집중은 재료 물성보다 형상에 의해 본질적으로 결정됨을 알 수 있다.")

subheading("3) 파라미터 실험: 구멍 크기(r) 변경")
body("선택 변수로 구멍 크기 r을 0.05부터 0.30까지 변화시키며 홀 가장자리의 최대 응력과 응력집중계수"
     "(SCF = σ_max / σ_nom)의 변화를 분석하였다. 공칭응력 σ_nom은 적용 far-field 응력 σ = E·Δ/H = 21 MPa로 "
     "정의하였다. 이 값은 작은 구멍(r=0.05)에서 유한요소 반력으로 산출한 평균응력(약 21.3 MPa)과 거의 일치하여 "
     "기준 응력으로서 타당하다. 최대 응력은 von Mises 응력의 최댓값을 사용하였다. 결과는 다음과 같다.")
make_table(["r (m)", "2r/W", "Max σ (von Mises) [MPa]", "σ_nom [MPa]", "SCF (=σ/σ_nom)"],
           [["0.05", "0.05", "55.4", "21", "2.64"],
            ["0.10", "0.10", "52.5", "21", "2.50"],
            ["0.15", "0.15", "51.5", "21", "2.45"],
            ["0.20", "0.20", "49.2", "21", "2.34"],
            ["0.25", "0.25", "44.8", "21", "2.13"],
            ["0.30", "0.30", "41.5", "21", "1.98"]])
two_images(os.path.join(BASE, "HW1_maxstress_vs_holesize.png"), "그림 3. 구멍 크기에 따른 최대 응력 변화",
           os.path.join(BASE, "HW1_SCF_vs_holesize.png"), "그림 4. 구멍 크기에 따른 SCF 변화")
body("[결과 해석] 최대 응력과 SCF 모두 구멍 크기가 커질수록 단조 감소하는 경향을 보였다(최대 응력 55.4→41.5 MPa, "
     "SCF 2.64→1.98). 이는 본 해석이 변위 구속 조건이기 때문이다. 즉 구멍이 커질수록 평판의 강성이 낮아져 같은 "
     "강제 변위 Δ를 주더라도 판이 전달하는 하중(반력)이 줄고, 그 결과 홀 가장자리의 첨두 응력 자체가 감소한다. "
     "한편 가장 작은 구멍(r=0.05)에서 SCF는 약 2.64로, 무한 평판·소형 원공의 고전 Kirsch 이론값 3에 비교적 "
     "근접하여 해석의 타당성을 보여 준다.")
body("[해석상 고려사항] SCF가 이론값 3보다 다소 낮게 나타난 주된 원인은 요소 종류와 메쉬 밀도에 있다. 본 해석에 "
     "사용된 CST(상수변형률 삼각형)는 요소당 응력이 일정하여, 응력 구배가 매우 급격한 홀 가장자리에서 첨두 응력을 "
     "과소평가한다. 따라서 홀 주변 메쉬를 더 조밀하게 하면 첨두 응력이 증가해 SCF가 이론값 3에 더 근접할 것으로 "
     "판단된다. 또한 본 시편은 높이 H가 짧고 구멍이 상·하 경계에 근접해 무한 스트립 가정에서 벗어나므로, 구멍이 "
     "커질수록 이론값과의 차이가 커지는 유한 효과도 함께 작용한다.")

# =====================================================================
# HOMEWORK #2
# =====================================================================
heading("Homework #2. 압입시험 데이터 기반 탄성계수(E) 예측 인공신경망", size=13, space_before=0)
body("수업에서 항복강도(YS)를 예측하던 동일 데이터셋(4,500개)을 사용하되, 목표변수(target)를 탄성계수 E로 "
     "변경하여 인공신경망(ANN)을 학습하였다. 입력변수(12개)는 Pmax, Stiffness2(제하 강성), Slope1~5, "
     "30~150 µm 변위에서의 하중값이다.")

subheading("1) 모델 구성 및 학습 과정")
bullet("모델: 다층 퍼셉트론(MLP) — Dense(128)-Dense(128)-Dropout(0.2)-Dense(64)-Dense(32)-Dense(1), 활성함수 ReLU, L2 정규화.")
bullet("손실함수 MAPE, 최적화기 Adam(lr=1e-3), 입력 Min-Max 정규화, 데이터 분할 70/15/15(train/val/test), EarlyStopping 적용.")
one_image(os.path.join(DS, "E_mape_curve.png"), "그림 4. Epoch에 따른 MAPE 변화 곡선 (Training/Validation)", w=11.0)
body("[학습 과정 해석] 그림 4의 MAPE 곡선은 학습이 진행되는 동안 모델의 거동을 단계적으로 보여 준다. "
     "①(초기 급감, 약 0~10 epoch) Training MAPE는 약 61%, Validation MAPE는 약 23%에서 출발하여 불과 "
     "십여 epoch 만에 수 % 수준으로 급격히 감소한다. 이는 신경망이 입력 변수와 탄성계수 사이의 가장 지배적인 "
     "관계(특히 제하 강성과 E의 비례 관계)를 매우 빠르게 학습함을 의미한다. ②(최적 일반화 시점) 검증 MAPE는 "
     "약 8 epoch에서 최저에 도달하여, 모델이 처음 보는 데이터에 대해 가장 잘 일반화되는 지점이 "
     "된다. ③(과적합 구간) 이후 Training MAPE는 약 1%까지 계속 감소하지만 Validation MAPE는 오히려 상승하여 "
     "약 15% 수준에서 정체(plateau)한다. 두 곡선이 크게 벌어진 채 1000 epoch까지 평탄하게 유지되는 것은 "
     "전형적인 과적합(overfitting)으로, 모델이 학습 데이터의 세부·잡음까지 외우기 시작하면서 일반화 성능이 "
     "더 이상 개선되지 않음을 보여 준다. ④(EarlyStopping의 역할) 검증 오차가 patience(1000 epoch) 동안 "
     "개선되지 않자 약 1000 epoch 부근에서 학습이 자동 종료되었고, restore_best_weights 설정에 따라 최종 "
     "모델은 과적합된 후반 상태가 아니라 ②의 최적 시점(epoch 8) 가중치로 복원되었다. 그 결과 테스트 성능이 "
     "MAPE 1.27%로 우수하게 유지될 수 있었다.")

subheading("2) 예측 성능 평가 및 예측 정확도 분석 (Test data)")
body("최적 모델(epoch 8의 복원된 가중치)을 학습에 사용하지 않은 테스트 데이터에 적용하여 예측 성능을 "
     "정량적으로 평가하였다. 결과는 다음과 같다.")
make_table(["지표", "MAPE", "MAE", "RMSE"],
           [["값", "1.27 %", "1.92 GPa", "2.68 GPa"]])
one_image(os.path.join(DS, "E_reference_vs_prediction.png"),
          "그림 5. 실제값(Reference) vs 예측값(Prediction) — Target: E", w=8.2)
body("[성능 해석] 탄성계수의 분포 범위가 60~240 GPa임을 고려하면, MAPE 1.27%와 MAE 약 1.92 GPa는 매우 "
     "높은 예측 정확도에 해당한다. RMSE(2.68 GPa)가 MAE보다 다소 큰 것은 일부 표본에서 상대적으로 큰 오차가 "
     "존재함을 의미하지만, 두 값의 차이가 크지 않아 예측을 크게 벗어나는 이상치(outlier)는 거의 없음을 알 수 "
     "있다.")
body("[예측 정확도 분석] 그림 5의 실제값–예측값 산점도는 거의 모든 점이 이상선(y=x) 위에 조밀하게 분포하여, "
     "낮은 탄성계수부터 높은 탄성계수까지 전 구간에서 고르게 정확함을 보여 준다. 다만 E가 큰 영역(약 220 GPa "
     "이상)에서 산포가 다소 커지는 경향이 관찰되는데, 이는 해당 구간의 학습 표본이 상대적으로 적거나 고강성 "
     "영역의 거동이 더 복잡하기 때문으로 해석된다. 참고로 동일 모델로 항복강도(YS)를 예측했을 때(MAPE 약 15%)"
     "보다 훨씬 우수한데, 이는 탄성계수가 제하 강성과 직접 연결되어 압입 데이터로부터 예측하기 용이한 물성이기 "
     "때문이다.")

subheading("3) 입력 변수 중요도 분석 (Permutation Importance & SHAP)")
body("어떤 입력 변수가 E 예측에 가장 큰 영향을 미치는지 Permutation Importance(변수를 무작위로 섞었을 때 오차 "
     "증가량)와 SHAP(예측 기여도) 두 가지 방법으로 분석하였다.")
two_images(os.path.join(DS, "E_permutation_importance.png"), "그림 6. Permutation Importance",
           os.path.join(DS, "E_shap_importance.png"), "그림 7. SHAP 변수 중요도")
subheading("결과 해석")
bullet("두 방법 모두 Stiffness2(제하 강성)가 압도적 1위로, 나머지 변수를 모두 합친 것보다 영향력이 훨씬 크다. (Permutation: 해당 변수 셔플 시 MAPE +50%p, SHAP: 평균 영향 ≈46 GPa로 2위 대비 ~15배)")
bullet("이는 압입역학적으로 타당하다. Oliver–Pharr 이론에서 탄성계수 Er은 제하 강성 S와 직접 비례한다(Er = (√π/2)·S/√A). 즉, 신경망이 데이터만으로 탄성계수–제하강성의 물리적 관계를 스스로 학습했음을 의미한다.")
bullet("보조적으로 Slope4·Slope5(제하 구간 기울기) 등이 미미하게 기여하는 것도 제하 거동과 관련되어 합리적이다.")

subheading("4) 학습 특성 및 분석 방법론")
bullet("정규화와 과적합 제어: 입력 Min-Max 정규화로 학습 안정성을 확보하고, Dropout(0.2)·L2 정규화·EarlyStopping으로 과적합을 제어하였다. 특히 EarlyStopping의 가중치 복원 기능 덕분에, 학습 후반의 과적합 경향에도 불구하고 최종 모델은 검증 성능이 가장 우수한 상태로 선택되었다.")
bullet("방법론 설명: Permutation Importance는 특정 입력 변수의 값만 무작위로 섞은 뒤 예측 오차(MAPE)가 얼마나 증가하는지를 측정하는 기법으로, 증가량이 클수록 중요한 변수다. SHAP은 게임이론의 Shapley 값을 이용해 각 예측에 대한 변수별 기여도를 정량화한다. 서로 다른 두 방법이 동일한 결론(Stiffness2 지배)을 제시하여 결과의 신뢰성이 높다.")
bullet("한계 및 향후 과제: 제하 강성과 일부 Slope·변위 변수는 물리적으로 상관관계가 있어 중요도가 특정 변수에 집중될 수 있다. 또한 MAPE 지표는 작은 E값에서 상대오차에 민감하다. 향후 변수 간 상관관계 분석, 하이퍼파라미터 최적화, 실측 데이터로의 확장 등을 통해 모델을 개선할 수 있다.")

subheading("5) 결론")
body("압입 데이터 기반 ANN은 탄성계수 E를 MAPE ≈1.2%의 높은 정확도로 예측하였으며, 변수 중요도 분석 결과 제하 "
     "강성이 결정적 변수로 나타났다. 이는 압입이론의 물리적 관계와 일치하여, 모델이 물리적으로 타당한 학습을 "
     "수행했음을 확인하였다.")

out = os.path.join(BASE, "과제보고서_HW1_HW2.docx")
doc.save(out)
print("Saved:", out)
